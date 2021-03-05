# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""

"""""""""""""""
IMPORTING PACKAGES
"""""""""""""""
import os
import math
import pandas as pd
import numpy as np
from scipy.io import loadmat
import matplotlib.pyplot as plt
import seaborn as sns

from pyeeg import bin_power, spectral_entropy
from docx import Document
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix, accuracy_score
from sklearn.ensemble import RandomForestClassifier

"""""""""""""""
FUNCTIONS
"""""""""""""""
# Remove zero-fil-in rows
def remove_zeros(data):
    new_data = data[~np.all(data == 0, axis = 1)]
    return new_data

# Calculate psis & spectral entropies
def get_psi_entropies(data):
    # Setting variables needed for for loop
    band = [0.5, 4, 7, 12, 30, 100]
    fs = 1024
    size = 1024
    
    columns = data.shape[1]
    rows = math.floor(data.shape[0] / size)
    alpha_psi_df = pd.DataFrame([])
    temp_data = pd.DataFrame([])
    
    # Creating for loop
    for x in range(columns):
        alpha_psis = []
        entropies = []
        temp_col_alpha = 'alpha_psi ' + str(x)
        temp_col_entropy = 'spectral_entropy ' + str(x)
        
        for y in range(rows):
            psis, power_ratios = bin_power(data.iloc[(y * size):((y + 1) * size), x], band, fs)
            alpha_psis.append(psis[2])
    
            spec_entropies = spectral_entropy(data.iloc[(y * size):((y + 1) * size), x], band, fs, power_ratios)
            entropies.append(spec_entropies)
        
        temp_data[temp_col_alpha] = alpha_psis
        temp_data[temp_col_entropy] = entropies
        
    alpha_psi_df = alpha_psi_df.append(temp_data, ignore_index = True)
    
    return alpha_psi_df

# Create state labels
def create_labels(pre_data, med_data, post_data): 
    pre_label = ["Pre"] * pre_data.shape[0]
    med_label = ["Med"] * med_data.shape[0]
    post_label = ["Post"] * post_data.shape[0]
    
    state_labels = pre_label + med_label + post_label
    
    return state_labels

# Combining datasets
def combine_datasets(data1, data2, data3):
    frames = [data1, data2, data3]
    new_df = pd.concat(frames, ignore_index = True)
    
    return new_df

# Remove co-linearity
def remove_colinearity(data, neg_threshold):
    corr_mat = data.corr()
    row = corr_mat.shape[0]
    column = corr_mat.shape[1]
    
    correlated_features = []
    
    for x in range(row): 
        for y in range(column):
            if x == y:
                break
            if corr_mat.iloc[x, y] > abs(neg_threshold) or corr_mat.iloc[x, y] < neg_threshold:
                correlated_features.append(corr_mat.columns[x])
                break

    return corr_mat, correlated_features

# Normalize data using min max scaler
def min_max(data): 
    scaler = MinMaxScaler()
    normalized_corr = scaler.fit_transform(data)
    normalied_corr_df = pd.DataFrame(normalized_corr)
    
    return normalied_corr_df

"""""""""""""""
SUBJECT 1 TASKS
"""""""""""""""
"""
Q1 - Load the data in each of the Matlab files into a data frame.
"""
# Setting current directory and getting files
os.chdir('/Users/umreenimam/Documents/Masters/Masters_Classes/CS_5310/chapter05/lab_data/Data/Subject1')
sub1_pre_file = 'Pre.mat'
sub1_med_file = 'Med.mat'
sub1_post_file = 'Post.mat'

# Loading data 
sub1_pre_data = loadmat(sub1_pre_file)
sub1_med_data = loadmat(sub1_med_file)
sub1_post_data = loadmat(sub1_post_file)

# Extracting data
sub1_pre_ = sub1_pre_data['data']
sub1_med_ = sub1_med_data['data']
sub1_post_ = sub1_post_data['data']

# Create dataframes and transpose them 
sub1_pre_t = pd.DataFrame(sub1_pre_.T)
sub1_med_t = pd.DataFrame(sub1_med_.T)
sub1_post_t = pd.DataFrame(sub1_post_.T)

"""
Q2 - Remove zero-fill-in rows from each data frame.
"""
# Calling remove_zeros function
sub1_pre = remove_zeros(sub1_pre_t)
sub1_med = remove_zeros(sub1_med_t)
sub1_post = remove_zeros(sub1_post_t)

"""
Q3 - Compute alpha PSIs and spectral entropies.
"""
# Calling alpha psi function on each dataframe
sub1_pre_df = get_psi_entropies(sub1_pre)
sub1_med_df = get_psi_entropies(sub1_med)
sub1_post_df = get_psi_entropies(sub1_post)

"""
Q4 & Q5 - Create list of brain state labels & combine dataframes into one.
"""
# Calling create_labels function
sub1_state_labels = create_labels(sub1_pre_df, sub1_med_df, sub1_post_df)


# Calling combine_datasets function
sub1_df = combine_datasets(sub1_pre_df, sub1_med_df, sub1_post_df)

"""
Q6-Q8 - Create a correlation coefficient matrix of 68 channels, remove co-linearity, 
and print new correlation matrix.
"""
# Running co-linearity function
corrleation_mat, cols_to_remove = remove_colinearity(sub1_df, -0.9)
sub1_df_remain = sub1_df.drop(columns = cols_to_remove, axis = 1)

# Correlation matrix of remaining data 
sub1_corr = sub1_df_remain.corr()

# Heatmaps of data before removing co-linearity & after removing 
# co-linearity
sns.set_theme(style = "white")

plt.figure(figsize=(8,8))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(corrleation_mat, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.1,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('sub1_fig1.png')
plt.show


plt.figure(figsize=(8,8))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(sub1_corr, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.5,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('sub1_fig2.png')
plt.show

"""
Q9 - Normalize data using min-max method.
"""
# Calling min_max function
sub1_normalized_corr_df = min_max(sub1_df_remain)

"""
Q10 - Split data into 80% training and 20% testing sets.
"""
X = sub1_normalized_corr_df
y = sub1_state_labels

# Splitting data into train and test models
sub1_X_train, sub1_X_test, sub1_y_train, sub1_y_test = train_test_split(
    X, 
    y, 
    random_state = None,
    test_size = 0.2,
    stratify = y)
     
"""
Q11 & Q12 - Train and test data using Random Forest Classifier.
"""
# Random Forest config
sub1_rf = RandomForestClassifier(
    n_estimators = 10, 
    max_depth = None, 
    min_samples_split = 2, 
    random_state = None)

# Training data
sub1_rf.fit(sub1_X_train, sub1_y_train)

# Testing data
sub1_pred = sub1_rf.predict(sub1_X_test)

"""
Q13 - Create confusion matrix and compute accuracy.
"""
# Confusion matrix creation
sub1_conf_mat = confusion_matrix(sub1_y_test, sub1_pred)

# Accuracy score
sub1_accuracy = accuracy_score(sub1_y_test, sub1_pred)
print('Accuracy Rate: {}%'.format(round(sub1_accuracy * 100, 1)))


"""
SUBJECT 2 TASKS
"""
"""
Q1 - Load the data in each of the Matlab files into a data frame.
"""
# Setting current directory and getting files
os.chdir('/Users/umreenimam/Documents/Masters/Masters_Classes/CS_5310/chapter05/lab_data/Data/Subject2')
sub2_pre_file = 'Pre.mat'
sub2_med_file = 'Med.mat'
sub2_post_file = 'Post.mat'

# Loading data 
sub2_pre_data = loadmat(sub2_pre_file)
sub2_med_data = loadmat(sub2_med_file)
sub2_post_data = loadmat(sub2_post_file)

# Extracting data
sub2_pre_ = sub2_pre_data['data']
sub2_med_ = sub2_med_data['data']
sub2_post_ = sub2_post_data['data']

# Create dataframes and transpose them 
sub2_pre_t = pd.DataFrame(sub2_pre_.T)
sub2_med_t = pd.DataFrame(sub2_med_.T)
sub2_post_t = pd.DataFrame(sub2_post_.T)

"""
Q2 - Remove zero-fill-in rows from each data frame.
"""
# Calling remove_zeros function
sub2_pre = remove_zeros(sub2_pre_t)
sub2_med = remove_zeros(sub2_med_t)
sub2_post = remove_zeros(sub2_post_t)

"""
Q3 - Compute alpha PSIs and spectral entropies.
"""
# Calling alpha psi function on each dataframe
sub2_pre_df = get_psi_entropies(sub2_pre)
sub2_med_df = get_psi_entropies(sub2_med)
sub2_post_df = get_psi_entropies(sub2_post)

"""
Q4 & Q5 - Create list of brain state labels & combine dataframes into one.
"""
# Calling create_labels function
sub2_state_labels = create_labels(sub2_pre_df, sub2_med_df, sub2_post_df)

# Calling combine_datasets function
sub2_df = combine_datasets(sub2_pre_df, sub2_med_df, sub2_post_df)

"""
Q6-Q8 - Create a correlation coefficient matrix of 68 channels, remove co-linearity, 
and print new correlation matrix.
"""
# Running co-linearity function
corrleation_mat, cols_to_remove = remove_colinearity(sub2_df, -0.9)
sub2_df_remain = sub2_df.drop(columns = cols_to_remove, axis = 1)

# Correlation matrix of remaining data 
sub2_corr = sub2_df_remain.corr()

# Heatmaps of data before removing co-linearity & after removing 
# co-linearity
sns.set_theme(style = "white")

plt.figure(figsize=(8,8))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(corrleation_mat, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.1,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('sub2_fig1.png')
plt.show


plt.figure(figsize=(8,8))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(sub2_corr, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.5,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('sub2_fig2.png')
plt.show

"""
Q9 - Normalize data using min-max method.
"""
# Calling min_max function
sub2_normalized_corr_df = min_max(sub2_df_remain)

"""
Q10 - Split data into 80% training and 20% testing sets.
"""
X = sub2_normalized_corr_df
y = sub2_state_labels

# Splitting data into train and test models
sub2_X_train, sub2_X_test, sub2_y_train, sub2_y_test = train_test_split(
    X, 
    y, 
    random_state = None,
    test_size = 0.2,
    stratify = y)
     
"""
Q11 & Q12 - Train and test data using Random Forest Classifier.
"""
# Random Forest config
sub2_rf = RandomForestClassifier(
    n_estimators = 10, 
    max_depth = None, 
    min_samples_split = 2, 
    random_state = None)

# Training data
sub2_rf.fit(sub2_X_train, sub2_y_train)

# Testing data
sub2_pred = sub2_rf.predict(sub2_X_test)

"""
Q13 - Create confusion matrix and compute accuracy.
"""
# Confusion matrix creation
sub2_conf_mat = confusion_matrix(sub2_y_test, sub2_pred)

# Accuracy score
sub2_accuracy = accuracy_score(sub2_y_test, sub2_pred)
print('Accuracy Rate: {}%'.format(round(sub2_accuracy * 100, 1)))

"""
SUBJECT 1 & 2 COMBINED
"""
"""
Q15 & Q16 - Combine subject 1 & 2 datasets. Repeat steps 6-13 for combined data.
"""
# Combine subject 1 & 2 dataframes into one
combined_frames = [sub1_df, sub2_df]
combined_df = pd.concat(combined_frames, ignore_index = True)

# Combine state labels
combined_labels = sub1_state_labels + sub2_state_labels

# Correlation matrix & removing co-linearity
combined_corr_mat, cols_to_remove = remove_colinearity(combined_df, -0.9)
combined_df_remain = combined_df.drop(columns = cols_to_remove, axis = 1)
combined_corr = combined_df_remain.corr()

# Heatmap of Correlation matrix before & after removing co-linearity
sns.set_theme(style = "white")

plt.figure(figsize=(8,8))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(combined_corr_mat, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.1,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Combined Data Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('combined_fig1.png')
plt.show


plt.figure(figsize=(8,8))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(combined_corr, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.5,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Combined Data Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('combined_fig2.png')
plt.show

# Normalize data using min_max function
combined_normalized = min_max(combined_df_remain)

# Split data into train & test 
X = combined_normalized
y = combined_labels

# Splitting data into train and test models
combined_X_train, combined_X_test, combined_y_train, combined_y_test = train_test_split(
    X, 
    y, 
    random_state = None,
    test_size = 0.2,
    stratify = y)

# Random Forest config
combined_rf = RandomForestClassifier(
    n_estimators = 10, 
    max_depth = None, 
    min_samples_split = 2, 
    random_state = None)

# Training data
combined_rf.fit(combined_X_train, combined_y_train)

# Testing data
combined_pred = combined_rf.predict(combined_X_test)

# Confusion matrix creation
combined_conf_mat = confusion_matrix(combined_y_test, combined_pred)

# Accuracy score
combined_accuracy = accuracy_score(combined_y_test, combined_pred)
print('Accuracy Rate: {}%'.format(round(combined_accuracy * 100, 1)))

"""
WORD DOC WITH LAB RESULTS
"""
doc = Document()

# SUBJECT 1 FINDINGS
doc.add_heading('Subject 1 Correlation Coefficient Matrix Before Removing Co-linearity', level = 1)

doc.add_picture('sub1_fig1.png')
doc.add_paragraph()
doc.add_page_break()

doc.add_heading('Subject 1 Correlation Coefficient Matrix After Removing Co-linearity', level = 1)

doc.add_picture('sub1_fig2.png')
doc.add_paragraph()

doc.add_heading('Confusion Matrix:', level = 1)

table = doc.add_table(rows = sub1_conf_mat.shape[0] + 1, cols = sub1_conf_mat.shape[1] + 1)
table.style = 'Medium Grid 3 Accent 3'

row = table.rows[0]
row.cells[1].text = 'Predicted Med'
row.cells[2].text = 'Predicted Post'
row.cells[3].text = 'Predicted Pre'

col = table.columns[0]
col.cells[1].text = 'Actual Med'
col.cells[2].text = 'Actual Post'
col.cells[3].text = 'Actual Pre'

for i in range(sub1_conf_mat.shape[0]):
    for j in range(sub1_conf_mat.shape[1]):
        table.cell(i + 1, j + 1).text = str(sub1_conf_mat[i, j])

doc.add_paragraph()
doc.add_paragraph('Subject 1 Accuracy Rate: {}%'.format(round(sub1_accuracy * 100, 1)))

# SUBJECT 2 FINDINGS
doc.add_heading('Subject 2 Correlation Coefficient Matrix Before Removing Co-linearity', level = 1)

doc.add_picture('sub2_fig1.png')
doc.add_paragraph()
doc.add_page_break()

doc.add_heading('Subject 2 Correlation Coefficient Matrix After Removing Co-linearity', level = 1)

doc.add_picture('sub2_fig2.png')
doc.add_paragraph()

doc.add_heading('Confusion Matrix:', level = 1)

table = doc.add_table(rows = sub2_conf_mat.shape[0] + 1, cols = sub2_conf_mat.shape[1] + 1)
table.style = 'Medium Grid 3 Accent 3'

row = table.rows[0]
row.cells[1].text = 'Predicted Med'
row.cells[2].text = 'Predicted Post'
row.cells[3].text = 'Predicted Pre'

col = table.columns[0]
col.cells[1].text = 'Actual Med'
col.cells[2].text = 'Actual Post'
col.cells[3].text = 'Actual Pre'

for i in range(sub2_conf_mat.shape[0]):
    for j in range(sub2_conf_mat.shape[1]):
        table.cell(i + 1, j + 1).text = str(sub2_conf_mat[i, j])

doc.add_paragraph()
doc.add_paragraph('Subject 2 Accuracy Rate: {}%'.format(round(sub2_accuracy * 100, 1)))

# COMBINED DATA FINDINGS
doc.add_heading('Combined Correlation Coefficient Matrix Before Removing Co-linearity', level = 1)

doc.add_picture('combined_fig1.png')
doc.add_paragraph()
doc.add_page_break()

doc.add_heading('Combined Correlation Coefficient Matrix After Removing Co-linearity', level = 1)

doc.add_picture('combined_fig2.png')
doc.add_paragraph()

doc.add_heading('Confusion Matrix:', level = 1)

table = doc.add_table(rows = combined_conf_mat.shape[0] + 1, cols = combined_conf_mat.shape[1] + 1)
table.style = 'Medium Grid 3 Accent 3'

row = table.rows[0]
row.cells[1].text = 'Predicted Med'
row.cells[2].text = 'Predicted Post'
row.cells[3].text = 'Predicted Pre'

col = table.columns[0]
col.cells[1].text = 'Actual Med'
col.cells[2].text = 'Actual Post'
col.cells[3].text = 'Actual Pre'

for i in range(combined_conf_mat.shape[0]):
    for j in range(combined_conf_mat.shape[1]):
        table.cell(i + 1, j + 1).text = str(combined_conf_mat[i, j])

doc.add_paragraph()
doc.add_paragraph('Combined Accuracy Rate: {}%'.format(round(combined_accuracy * 100, 1)))

doc.save('Chapter05-Lab-UI2.docx')