# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""

"""""""""""""""
IMPORTING PACKAGES
"""""""""""""""
import os
import math
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import wittgenstein as lw

from pyeeg import bin_power, spectral_entropy
from docx import Document
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix, accuracy_score, classification_report

"""""""""""""""
FUNCTIONS
"""""""""""""""
# Function to read and load data
def load_read_data(data1, data2, data3):
    subject_data1 = pd.read_csv(data1)
    subject_data2 = pd.read_csv(data2)
    subject_data3 = pd.read_csv(data3)
    
    return subject_data1, subject_data2, subject_data3

# Calculate psis & spectral entropies
def get_psi_entropies(data):
    # Setting variables needed for for loop
    band = [0.5, 4, 7, 12, 30, 100]
    fs = 1024
    size = 1024
    
    columns = data.shape[1]
    rows = math.floor(data.shape[0] / size)
    alpha_psi_df = pd.DataFrame([])
    temp_data = pd.DataFrame([])
    
    # Creating for loop
    for x in range(columns):
        alpha_psis = []
        entropies = []
        temp_col_alpha = 'alpha_psi ' + str(x)
        temp_col_entropy = 'spectral_entropy ' + str(x)
        
        for y in range(rows):
            psis, power_ratios = bin_power(data.iloc[(y * size):((y + 1) * size), x], band, fs)
            alpha_psis.append(psis[2])
    
            spec_entropies = spectral_entropy(data.iloc[(y * size):((y + 1) * size), x], band, fs, power_ratios)
            entropies.append(spec_entropies)
        
        temp_data[temp_col_alpha] = alpha_psis
        temp_data[temp_col_entropy] = entropies
        
    alpha_psi_df = alpha_psi_df.append(temp_data, ignore_index = True)
    
    return alpha_psi_df

# Create state labels
def create_labels(pre_data, med_data, post_data): 
    pre_label = ["Pre"] * pre_data.shape[0]
    med_label = ["Med"] * med_data.shape[0]
    post_label = ["Post"] * post_data.shape[0]
    
    state_labels = pre_label + med_label + post_label
    
    return state_labels

# Combining datasets
def combine_datasets(data1, data2, data3):
    frames = [data1, data2, data3]
    new_df = pd.concat(frames, ignore_index = True)
    
    return new_df

# Remove co-linearity
def remove_colinearity(data, neg_threshold):
    corr_mat = data.corr()
    row = corr_mat.shape[0]
    column = corr_mat.shape[1]
    
    correlated_features = []
    
    for x in range(row): 
        for y in range(column):
            if x == y:
                break
            if corr_mat.iloc[x, y] > abs(neg_threshold) or corr_mat.iloc[x, y] < neg_threshold:
                correlated_features.append(corr_mat.columns[x])
                break

    return corr_mat, correlated_features

# Normalize data using min max scaler
def min_max(data): 
    scaler = MinMaxScaler()
    normalized_corr = scaler.fit_transform(data)
    normalied_corr_df = pd.DataFrame(normalized_corr)
    
    return normalied_corr_df

# Creation of heatmap function
def create_heatmap(data, figure_name):
    sns.set_theme(style = "white")
    
    plt.figure(figsize=(12, 12))
    color_map = sns.diverging_palette(230, 20, as_cmap = True)
    sns.heatmap(data, annot = False, cmap = color_map, vmax = 1, 
                center = 0, square = True, linewidths = 0.1,
                cbar_kws = {"shrink": 0.75})
    plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
    plt.xlabel('Column Number from the Data Frame', fontsize = 12)
    plt.ylabel('Column Number from the Data Frame', fontsize = 12)
    plt.savefig(figure_name)
    plt.show

# Function to get highest probability of state lables from trained data
states = ['Pre', 'Med', 'Post']

def state_res(pre_prob, med_prob, post_prob):
    if abs(pre_prob[0]) < 0.0001 and abs(pre_prob[1]) < 0.0001:
        pre_prob[1] = 1.0
    if abs(med_prob[0]) < 0.0001 and abs(med_prob[1]) < 0.0001:
        med_prob[1] = 1.0
    if abs(post_prob[0]) < 0.0001 and abs(post_prob[1]) < 0.0001:
        post_prob[1] = 1.0
        
    res = [pre_prob[1], med_prob[1], post_prob[1]]
    max_value = max(res)
    max_index = res.index(max_value)
    
    return states[max_index]

def data_precision_recall(conf_mat):
    true_pos = np.diag(conf_mat)
    precision = np.mean(true_pos / np.sum(conf_mat, axis = 0))
    recall = np.mean(true_pos / np.sum(conf_mat, axis = 1))
    
    return precision, recall

"""""""""""""""
SUBJECT 1 TASKS
"""""""""""""""
"""
Q1 - Load the data in each of the csv files into a data frame.
"""
# Setting current directory and getting files
os.chdir('/Users/umreenimam/Documents/Masters/Masters_Classes/CS_5310/chapter05/lab_data/Data2/VR-EEG-Analysis')
sub1_pre_file = 'Subject1/Pre.csv'
sub1_med_file = 'Subject1/Med.csv'
sub1_post_file = 'Subject1/Post.csv'

# Calling function to read and load data into DataFrame
sub1_pre_df, sub1_med_df, sub1_post_df = load_read_data(
    sub1_pre_file,
    sub1_med_file,
    sub1_post_file)

"""
Q2 - Compute alpha PSIs and spectral entropies.
"""
# Calling alpha psi function for each dataframe
sub1_pre_alpha = get_psi_entropies(sub1_pre_df)
sub1_med_alpha = get_psi_entropies(sub1_med_df)
sub1_post_alpha = get_psi_entropies(sub1_post_df)

"""
Q3 & Q4 - Create list of brain state labels & combine dataframes into one.
"""
# Calling create_labels function
sub1_state_labels = create_labels(
    sub1_pre_alpha, 
    sub1_med_alpha, 
    sub1_post_alpha)


# Calling combine_datasets function
sub1_df = combine_datasets(
    sub1_pre_alpha, 
    sub1_med_alpha, 
    sub1_post_alpha)

"""
Q5-Q7 - Create a correlation coefficient matrix of 68 channels, 
remove co-linearity, and print new correlation matrix.
"""
# Calling co-linearity function
correlation_mat, cols_to_remove = remove_colinearity(sub1_df, -0.9)

# Remove co-linear columns 
sub1_remain_df = sub1_df.drop(columns = cols_to_remove, axis = 1)

# Correlation matrix of remaining data
sub1_corr = sub1_remain_df.corr()

# Heatmaps of data before removing co-linearity & after removing 
# co-linearity    
corr_mat_map = create_heatmap(correlation_mat, 'sub1_fig1.png')
corr_remain_map = create_heatmap(sub1_corr, 'sub1_fig2.png')

"""
Q8 - Normalize data using min-max method.
"""
# Calling min_max method 
sub1_normalized = min_max(sub1_remain_df)

"""
Q9 - Split data into 80% training and 20% testing sets.
"""
X = sub1_normalized
y = sub1_state_labels

# Splitting data
sub1_X_train, sub1_X_test, sub1_y_train, sub1_y_test = train_test_split(
    X,
    y,
    test_size = 0.2,
    random_state = 42,
    stratify = sub1_state_labels)

"""
Q10 - Train training dataset using "RIPPER" model.
"""
# Need to train with each brain state being the positive class
# Passing random state of different values for each data model
pre_data_model = lw.RIPPER(random_state = 42)
med_data_model = lw.RIPPER(random_state = 36)
post_data_model = lw.RIPPER(random_state = 28)

# Pre train
pre_data_model.fit(
    sub1_X_train, 
    sub1_y_train, 
    pos_class = 'Pre')

# Med train
med_data_model.fit(
    sub1_X_train, 
    sub1_y_train, 
    pos_class = 'Med')

# Post train
post_data_model.fit(
    sub1_X_train, 
    sub1_y_train, 
    pos_class = 'Post')

"""
Q11 - Print out the rule set in the model.
"""
print('Pre Data ruleset: \n')
pre_data_model.ruleset_.out_pretty()

print('Med Data rulestet: \n')
med_data_model.ruleset_.out_pretty()

print('Post Data ruleset: \n')
post_data_model.ruleset_.out_pretty()

"""
Q12 - Test the classifier using the testing dataset.
"""
# Testing classifier using predict function
sub1_pre_prediction = pre_data_model.predict(sub1_X_test)
sub1_med_prediction = med_data_model.predict(sub1_X_test)
sub1_post_prediction = post_data_model.predict(sub1_X_test)

# Need to use predict_proba function  
sub1_pre_proba = pre_data_model.predict_proba(sub1_X_test)
sub1_med_proba = med_data_model.predict_proba(sub1_X_test)
sub1_post_proba = post_data_model.predict_proba(sub1_X_test)

"""
Q13 & Q14 - Create a confusion matrix & print out model performance. 
Compute the accuracy, precision, & recall of the prediction.
"""
# Calling state_results function
sub1_pred1 = [state_res(
    sub1_pre_proba[i], 
    sub1_med_proba[i], 
    sub1_post_proba[i]) for i in range(len(sub1_pre_proba))]

# Using loop results in confusion_matrix function
sub1_confusion = confusion_matrix(
    sub1_y_test, 
    sub1_pred1)
print(f'Subject 1 Confusion Matrix: \n {sub1_confusion}')

# Print out model performance using classification report
sub1_report = classification_report(
    sub1_y_test, 
    sub1_pred1)
print(f'Subject 1 Classification Report: \n {sub1_report}')

# Compute accuracy, precision, and recall prediction
# Accuracy
sub1_accuracy = accuracy_score(sub1_y_test, sub1_pred1)
print('Subject 1 Accuracy Rate: {}%'.format(round(sub1_accuracy * 100, 1)))

# Precision and recall
sub1_precision, sub1_recall = data_precision_recall(sub1_confusion)
print('Subject 1 Precision: {}%'.format(round(sub1_precision * 100, 1)))
print('Subject 1 Recall: {}%'.format(round(sub1_recall * 100, 1)))

"""""""""""""""
SUBJECT 2 TASKS
"""""""""""""""
"""
Q1 - Load the data in each of the csv files into a data frame.
"""
# Setting current directory and getting files
os.chdir('/Users/umreenimam/Documents/Masters/Masters_Classes/CS_5310/chapter05/lab_data/Data2/VR-EEG-Analysis')
sub2_pre_file = 'Subject2/Pre.csv'
sub2_med_file = 'Subject2/Med.csv'
sub2_post_file = 'Subject2/Post.csv'

# Calling function to read and load data into DataFrame
sub2_pre_df, sub2_med_df, sub2_post_df = load_read_data(
    sub2_pre_file,
    sub2_med_file,
    sub2_post_file)

"""
Q2 - Compute alpha PSIs and spectral entropies.
"""
# Calling alpha psi function for each dataframe
sub2_pre_alpha = get_psi_entropies(sub2_pre_df)
sub2_med_alpha = get_psi_entropies(sub2_med_df)
sub2_post_alpha = get_psi_entropies(sub2_post_df)

"""
Q3 & Q4 - Create list of brain state labels & combine dataframes into one.
"""
# Calling create_labels function
sub2_state_labels = create_labels(
    sub2_pre_alpha, 
    sub2_med_alpha, 
    sub2_post_alpha)


# Calling combine_datasets function
sub2_df = combine_datasets(
    sub2_pre_alpha, 
    sub2_med_alpha, 
    sub2_post_alpha)

"""
Q5-Q7 - Create a correlation coefficient matrix of 68 channels, 
remove co-linearity, and print new correlation matrix.
"""
# Calling co-linearity function
correlation_mat, cols_to_remove = remove_colinearity(sub2_df, -0.9)

# Remove co-linear columns 
sub2_remain_df = sub2_df.drop(columns = cols_to_remove, axis = 1)

# Correlation matrix of remaining data
sub2_corr = sub2_remain_df.corr()

# Heatmaps of data before removing co-linearity & after removing 
# co-linearity
sub2_corr_map = create_heatmap(correlation_mat, 'sub2_fig1.png')
sub2_corr_remain_map = create_heatmap(sub2_corr, 'sub2_fig2.png')

"""
Q8 - Normalize data using min-max method.
"""
# Calling min_max method 
sub2_normalized = min_max(sub2_remain_df)

"""
Q9 - Split data into 80% training and 20% testing sets.
"""
X = sub2_normalized
y = sub2_state_labels

# Splitting data
sub2_X_train, sub2_X_test, sub2_y_train, sub2_y_test = train_test_split(
    X,
    y,
    test_size = 0.2,
    random_state = 42,
    stratify = sub2_state_labels)

"""
Q10 - Train training dataset using "RIPPER" model.
"""
# Need to train with each brain state being the positive class
# Passing random state of different values for model
pre_data_model2 = lw.RIPPER(random_state = 42)
med_data_model2 = lw.RIPPER(random_state = 36)
post_data_model2 = lw.RIPPER(random_state = 28)

# Pre train
pre_data_model2.fit(
    sub2_X_train, 
    sub2_y_train, 
    pos_class = 'Pre')

# Med train
med_data_model2.fit(
    sub2_X_train, 
    sub2_y_train, 
    pos_class = 'Med')

# Post train
post_data_model2.fit(
    sub2_X_train, 
    sub2_y_train, 
    pos_class = 'Post')

"""
Q11 - Print out the rule set in the model.
"""
print('Pre Data ruleset: \n')
pre_data_model2.ruleset_.out_pretty()

print('Med Data ruleset: \n')
med_data_model2.ruleset_.out_pretty()

print('Post Data ruleset: \n')
post_data_model2.ruleset_.out_pretty()

"""
Q12 - Test the classifier using the testing dataset.
"""
sub2_pre_prediction = pre_data_model2.predict(sub2_X_test)
sub2_med_prediction = med_data_model2.predict(sub2_X_test)
sub2_post_prediction = post_data_model2.predict(sub2_X_test)

# Need to use predict_proba function  
sub2_pre_proba = pre_data_model2.predict_proba(sub2_X_test)
sub2_med_proba = med_data_model2.predict_proba(sub2_X_test)
sub2_post_proba = post_data_model2.predict_proba(sub2_X_test)

"""
Q13 & Q14 - Create a confusion matrix & print out model performance. 
Compute the accuracy, precision, & recall of the prediction.
"""
# Calling state_results function
sub2_pred = [state_res(
    sub2_pre_proba[i], 
    sub2_med_proba[i], 
    sub2_post_proba[i]) for i in range(len(sub2_pre_proba))]

# Using loop results in confusion_matrix function
sub2_confusion = confusion_matrix(
    sub2_y_test, 
    sub2_pred)
print(f'Subject 2 Confustion Matrix: \n {sub2_confusion}')

# Print out model performance using classification report
sub2_report = classification_report(
    sub2_y_test,
    sub2_pred)
print(f'Subject 2 Classification Report: \n {sub2_report}')

# Compute accuracy, precision, and recall prediction
# Accuracy
sub2_accuracy = accuracy_score(sub2_y_test, sub2_pred)
print('Subject 2 Accuracy Rate: {}%'.format(round(sub2_accuracy * 100, 1)))

# Precision and recall
sub2_precision, sub2_recall = data_precision_recall(sub2_confusion)
print('Subject 2 Precision: {}%'.format(round(sub2_precision * 100, 1)))
print('Subject 2 Recall: {}%'.format(round(sub2_recall * 100, 1)))

"""""""""""""""
COMBINED SUBJECTS TASKS
"""""""""""""""
"""
Q1 - Combine subjects 1 & 2 data sets.
"""
# Combine subject 1 & 2 dataframes into one
combined_frames = [sub1_df, sub2_df]
combined_df = pd.concat(combined_frames, ignore_index = True)

# Combine state labels
combined_labels = sub1_state_labels + sub2_state_labels

"""
Q5-Q7 - Create a correlation coefficient matrix of 68 channels, 
remove co-linearity, and print new correlation matrix.
"""
# Correlation matrix & removing co-linearity
combined_corr_mat, cols_to_remove = remove_colinearity(combined_df, -0.9)
combined_df_remain = combined_df.drop(columns = cols_to_remove, axis = 1)
combined_corr = combined_df_remain.corr()

# Heatmaps of data before removing co-linearity & after removing 
# co-linearity
combined_corr_map = create_heatmap(combined_corr_mat, 'comb_fig1.png')
combined_corr_remain_map = create_heatmap(combined_corr, 'comb_fig2.png')

"""
Q8 - Normalize data using min-max method.
"""
# Calling min_max method 
combined_normalized = min_max(combined_df_remain)

"""
Q9 - Split data into 80% training and 20% testing sets.
"""
X = combined_normalized
y = combined_labels

# Splitting data
combined_X_train, combined_X_test, combined_y_train, combined_y_test = train_test_split(
    X,
    y,
    test_size = 0.2,
    random_state = 42,
    stratify = combined_labels)

"""
Q10 - Train training dataset using "RIPPER" model.
"""
# Need to train with each brain state being the positive class
# Passing random state of different values for model
pre_data_combined = lw.RIPPER(random_state = 42)
med_data_combined = lw.RIPPER(random_state = 36)
post_data_combined = lw.RIPPER(random_state = 28)

# Pre train
pre_data_combined.fit(
    combined_X_train, 
    combined_y_train, 
    pos_class = 'Pre')

# Med train
med_data_combined.fit(
    combined_X_train, 
    combined_y_train, 
    pos_class = 'Med')

# Post train
post_data_combined.fit(
    combined_X_train, 
    combined_y_train, 
    pos_class = 'Post')

"""
Q11 - Print out the rule set in the model.
"""
print('Pre data combined ruleset: \n')
pre_data_combined.ruleset_.out_pretty()

print('Med data combined ruleset: \n')
med_data_combined.ruleset_.out_pretty()

print('Post data combined ruleset: \n')
post_data_combined.ruleset_.out_pretty()

"""
Q12 - Test the classifier using the testing dataset.
"""
combined_pre_prediction = pre_data_combined.predict(combined_X_test)
combined_med_prediction = med_data_combined.predict(combined_X_test)
combined_post_prediction = post_data_combined.predict(combined_X_test)

# Need to use predict_proba function  
combined_pre_proba = pre_data_combined.predict_proba(combined_X_test)
combined_med_proba = med_data_combined.predict_proba(combined_X_test)
combined_post_proba = post_data_combined.predict_proba(combined_X_test)

"""
Q13 & Q14 - Create a confusion matrix & print out model performance. 
Compute the accuracy, precision, & recall of the prediction.
"""
# Calling state_results function
combined_pred = [state_res(
    combined_pre_proba[i], 
    combined_med_proba[i], 
    combined_post_proba[i]) for i in range(len(combined_pre_proba))]

# Using loop results in confusion_matrix function
combined_confusion = confusion_matrix(
    combined_y_test, 
    combined_pred)
print(f'Combined Confustion Matrix: \n {combined_confusion}')

# Print out model performance using classification report
combined_report = classification_report(
    combined_y_test,
    combined_pred)
print(f'Combined Classification Report: \n {combined_report}')

# Compute accuracy, precision, and recall prediction
# Accuracy
combined_accuracy = accuracy_score(combined_y_test, combined_pred)
print('Combined Accuracy Rate: {}%'.format(round(combined_accuracy * 100, 1)))

# Precision and recall
combined_precision, combined_recall = data_precision_recall(combined_confusion)
print('Combined Precision: {}%'.format(round(combined_precision * 100, 1)))
print('Combined Recall: {}%'.format(round(combined_recall * 100, 1)))

"""
WORD DOC WITH LAB RESULTS
"""
doc = Document()

# SUBJECT 1 FINDINGS
doc.add_heading('Subject 1 Correlation Coefficient Matrix Before Removing Co-linearity', level = 1)

doc.add_picture('sub1_fig1.png')
doc.add_paragraph()
doc.add_page_break()

doc.add_heading('Subject 1 Correlation Coefficient Matrix After Removing Co-linearity', level = 1)

doc.add_picture('sub1_fig2.png')
doc.add_paragraph()

doc.add_heading('Confusion Matrix:', level = 1)

table = doc.add_table(rows = sub1_confusion.shape[0] + 1, cols = sub1_confusion.shape[1] + 1)
table.style = 'Medium Grid 3 Accent 3'

row = table.rows[0]
row.cells[1].text = 'Predicted Med'
row.cells[2].text = 'Predicted Post'
row.cells[3].text = 'Predicted Pre'

col = table.columns[0]
col.cells[1].text = 'Actual Med'
col.cells[2].text = 'Actual Post'
col.cells[3].text = 'Actual Pre'

for i in range(sub1_confusion.shape[0]):
    for j in range(sub1_confusion.shape[1]):
        table.cell(i + 1, j + 1).text = str(sub1_confusion[i, j])

doc.add_paragraph()
doc.add_paragraph('Subject 1 Accuracy Rate: {}%'.format(round(sub1_accuracy * 100, 1)))
doc.add_paragraph('Subject 1 Precision: {}%'.format(round(sub1_precision * 100, 1)))
doc.add_paragraph('Subject 1 Recall: {}%'.format(round(sub1_recall * 100, 1)))

# SUBJECT 2 FINDINGS
doc.add_heading('Subject 2 Correlation Coefficient Matrix Before Removing Co-linearity', level = 1)

doc.add_picture('sub2_fig1.png')
doc.add_paragraph()
doc.add_page_break()

doc.add_heading('Subject 2 Correlation Coefficient Matrix After Removing Co-linearity', level = 1)

doc.add_picture('sub2_fig2.png')
doc.add_paragraph()

doc.add_heading('Confusion Matrix:', level = 1)

table = doc.add_table(rows = sub2_confusion.shape[0] + 1, cols = sub2_confusion.shape[1] + 1)
table.style = 'Medium Grid 3 Accent 3'

row = table.rows[0]
row.cells[1].text = 'Predicted Med'
row.cells[2].text = 'Predicted Post'
row.cells[3].text = 'Predicted Pre'

col = table.columns[0]
col.cells[1].text = 'Actual Med'
col.cells[2].text = 'Actual Post'
col.cells[3].text = 'Actual Pre'

for i in range(sub2_confusion.shape[0]):
    for j in range(sub2_confusion.shape[1]):
        table.cell(i + 1, j + 1).text = str(sub2_confusion[i, j])

doc.add_paragraph()
doc.add_paragraph('Subject 2 Accuracy Rate: {}%'.format(round(sub2_accuracy * 100, 1)))
doc.add_paragraph('Subject 2 Precision: {}%'.format(round(sub2_precision * 100, 1)))
doc.add_paragraph('Subject 2 Recall: {}%'.format(round(sub2_recall * 100, 1)))

# COMBINED DATA FINDINGS
doc.add_heading('Combined Correlation Coefficient Matrix Before Removing Co-linearity', level = 1)

doc.add_picture('comb_fig1.png')
doc.add_paragraph()
doc.add_page_break()

doc.add_heading('Combined Correlation Coefficient Matrix After Removing Co-linearity', level = 1)

doc.add_picture('comb_fig2.png')
doc.add_paragraph()

doc.add_heading('Confusion Matrix:', level = 1)

table = doc.add_table(rows = combined_confusion.shape[0] + 1, cols = combined_confusion.shape[1] + 1)
table.style = 'Medium Grid 3 Accent 3'

row = table.rows[0]
row.cells[1].text = 'Predicted Med'
row.cells[2].text = 'Predicted Post'
row.cells[3].text = 'Predicted Pre'

col = table.columns[0]
col.cells[1].text = 'Actual Med'
col.cells[2].text = 'Actual Post'
col.cells[3].text = 'Actual Pre'

for i in range(combined_confusion.shape[0]):
    for j in range(combined_confusion.shape[1]):
        table.cell(i + 1, j + 1).text = str(combined_confusion[i, j])

doc.add_paragraph()
doc.add_paragraph('Combined Accuracy Rate: {}%'.format(round(combined_accuracy * 100, 1)))
doc.add_paragraph('Combined Precision: {}%'.format(round(combined_precision * 100, 1)))
doc.add_paragraph('Combined Recall: {}%'.format(round(combined_recall * 100, 1)))

doc.save('Chapter05-Part02-Lab.docx')